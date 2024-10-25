import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn

# Function to parse Nmap XML
def parse_nmap_xml(xml_file):
    tree = ET.parse(xml_file)
    root = tree.getroot()
    scan_data = []
    
    for host in root.findall('host'):
        host_data = {}
        address = host.find('address').attrib.get('addr', 'N/A')
        host_data['address'] = address
        status = host.find('status').attrib.get('state', 'N/A')
        host_data['status'] = status

        # Safely retrieve the hostname if it exists
        hostname_elem = host.find('hostnames/hostname')
        hostname = hostname_elem.attrib.get('name', 'N/A') if hostname_elem is not None else 'N/A'
        host_data['hostname'] = hostname
        
        ports = []
        for port in host.findall('./ports/port'):
            port_data = {
                'portid': port.attrib.get('portid', 'N/A'),
                'protocol': port.attrib.get('protocol', 'N/A'),
                'service': port.find('./service').attrib.get('name', 'N/A'),
                'product': port.find('./service').attrib.get('product', 'N/A') if port.find('./service') is not None else 'N/A',
                'version': port.find('./service').attrib.get('version', 'N/A') if port.find('./service') is not None else 'N/A',
                'extrainfo': port.find('./service').attrib.get('extrainfo', 'N/A') if port.find('./service') is not None else 'N/A',
                'method': port.find('./service').attrib.get('method', 'N/A') if port.find('./service') is not None else 'N/A',
                'state': port.find('./state').attrib.get('state', 'N/A'),
                'fingerprint': None,  # Initialize fingerprint as None
                'info': ''
            }

            # Extract script info for "Info:" row if available
            script_output = ''
            for script in port.findall('script'):
                script_id = script.attrib.get('id', 'N/A')
                script_output += f"{script_id}: {script.attrib.get('output', 'N/A')}\n"

                # If it's a fingerprint script, set it as fingerprint and skip adding to info
                if script_id == 'fingerprint-strings':
                    port_data['fingerprint'] = script.attrib.get('output', 'N/A')
                else:
                    script_output = script_output.strip()

            if script_output:
                port_data['info'] = script_output

            # Extract vulnerability info if available
            vulnerabilities = []
            for script in port.findall('script'):
                if script.attrib.get('id').startswith("vulners"):
                    vuln_data = {
                        'id': script.attrib.get('id', 'N/A'),
                        'output': script.attrib.get('output', 'N/A'),
                    }
                    # Parse CVE and CVSS score
                    cve_elems = script.findall('.//elem')
                    for elem in cve_elems:
                        if 'CVE' in elem.text:
                            vuln_data['cve'] = elem.text
                        if elem.tag == 'cvss':
                            vuln_data['cvss'] = elem.text
                    vulnerabilities.append(vuln_data)
            port_data['vulnerabilities'] = vulnerabilities
            ports.append(port_data)

        host_data['ports'] = ports
        scan_data.append(host_data)

    return scan_data


# Function to set font and style for cell text
def set_cell_text(cell, text, font_size=10, bold=False, color=RGBColor(0, 0, 0)):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color

# Function to add table borders
def add_borders(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Thinner borders
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    
    tbl.tblPr.append(tblBorders)

# Function to create the Word report
def create_nmap_report(scan_data, output_doc):
    doc = Document()

    # Title of the report
    doc.add_heading('Nmap Security Scan Report', 0)

    for host in scan_data:
        # Host and status table with notes spanning both columns
        table = doc.add_table(rows=2, cols=3)
        table.autofit = True
        add_borders(table)

        # Set the text for host address and status with proper styling
        set_cell_text(table.cell(0, 0), f"Host: {host['address']}", font_size=12, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))  # Blue header for host
        set_cell_text(table.cell(0, 1), f"Status: {host['status']}", font_size=12, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))  # Blue header for status
        set_cell_text(table.cell(0, 2), f"Hostname(PTR): {host['hostname']}", font_size=12, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))  # Blue header for hostname

        # Notes row, spans across both columns
        notes_cell = table.cell(1, 0)
        notes_cell.merge(table.cell(1, 1))
        notes_cell.merge(table.cell(1, 2))
        set_cell_text(notes_cell, 'Notes:', font_size=10, bold=True)
        notes_cell.paragraphs[0].add_run("\n\n\n")  # Add extra space for writing notes

        # Apply shading to the host table, but make the Notes cell background white
        for row in table.rows:
            for cell in row.cells:
                shading = OxmlElement("w:shd")
                if cell is notes_cell:
                    shading.set(qn("w:fill"), "FFFFFF")  # White for the notes cell
                else:
                    shading.set(qn("w:fill"), "DCE6F1")  # Light blue for other cells
                cell._element.get_or_add_tcPr().append(shading)

        doc.add_paragraph("")  # Add a space between host table and first port

        if not host['ports']:
            doc.add_paragraph("No open ports found.")
            continue

        # Port and Vulnerability Table
        for port in host['ports']:
            # Add a row with the port details spanning the full table width
            port_row = doc.add_table(rows=1, cols=1)
            add_borders(port_row)
            port_cell = port_row.rows[0].cells[0]
            set_cell_text(port_cell, f"Port {port['portid']} ({port['protocol']})", font_size=12, bold=True, color=RGBColor(0x4F, 0x81, 0xBD))  # Port header in bold and blue
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9E1F2")  # Light blue for the port row
            port_cell._element.get_or_add_tcPr().append(shading)

            # Service, product, method, and vulnerability info below the port header
            port_info_table = doc.add_table(rows=2, cols=6)
            add_borders(port_info_table)
            hdr_cells = port_info_table.rows[0].cells
            set_cell_text(hdr_cells[0], 'Service', font_size=10, bold=True)
            set_cell_text(hdr_cells[1], 'Product', font_size=10, bold=True)
            set_cell_text(hdr_cells[2], 'Version', font_size=10, bold=True)
            set_cell_text(hdr_cells[3], 'Info', font_size=10, bold=True)
            set_cell_text(hdr_cells[4], 'Method', font_size=10, bold=True)
            set_cell_text(hdr_cells[5], 'Vulnerabilities', font_size=10, bold=True)

            row_cells = port_info_table.rows[1].cells
            set_cell_text(row_cells[0], port.get('service', 'N/A'), font_size=10)
            set_cell_text(row_cells[1], port.get('product', 'N/A'), font_size=10)
            set_cell_text(row_cells[2], port.get('version', 'N/A'), font_size=10)
            set_cell_text(row_cells[3], port.get('extrainfo', 'N/A'), font_size=10)
            set_cell_text(row_cells[4], port.get('method', 'N/A'), font_size=10)
            set_cell_text(row_cells[5], 'Yes' if port['vulnerabilities'] else 'No', font_size=10)

            # Add fingerprint row spanning the full table width (only if fingerprint exists)
            if port['fingerprint']:
                fingerprint_row = port_info_table.add_row()
                fingerprint_cell = fingerprint_row.cells[0]
                fingerprint_cell.merge(fingerprint_row.cells[1])
                fingerprint_cell.merge(fingerprint_row.cells[2])
                fingerprint_cell.merge(fingerprint_row.cells[3])
                fingerprint_cell.merge(fingerprint_row.cells[4])
                fingerprint_cell.merge(fingerprint_row.cells[5])
                set_cell_text(fingerprint_cell, f"Fingerprint: {port['fingerprint']}", font_size=10)
                shading_fingerprint = OxmlElement("w:shd")
                shading_fingerprint.set(qn("w:fill"), "F2F2F2")  # Light grey for fingerprint data
                fingerprint_cell._element.get_or_add_tcPr().append(shading_fingerprint)

            # Add script info row if available and not duplicating fingerprint
            if port['info'] and (port['fingerprint'] is None or port['fingerprint'] not in port['info']):  # Ensure fingerprint data isn't duplicated
                info_row = port_info_table.add_row()
                info_cell = info_row.cells[0]
                info_cell.merge(info_row.cells[1])
                info_cell.merge(info_row.cells[2])
                info_cell.merge(info_row.cells[3])
                info_cell.merge(info_row.cells[4])
                info_cell.merge(info_row.cells[5])
                set_cell_text(info_cell, f"Info: {port['info']}", font_size=10)
                shading_info = OxmlElement("w:shd")
                shading_info.set(qn("w:fill"), "F5F5F5")  # Light grey for info row
                info_cell._element.get_or_add_tcPr().append(shading_info)

            # Add a space between ports
            doc.add_paragraph("")

    # Save the document
    doc.save(output_doc)


# Main function to handle CLI arguments
def main():
    import argparse
    parser = argparse.ArgumentParser(description="Nmap XML to Word Report Converter")
    parser.add_argument("xml_file", help="Path to the Nmap XML file")
    parser.add_argument("output_doc", help="Path for the output Word document")

    args = parser.parse_args()

    # Parse the XML file and create the report
    scan_data = parse_nmap_xml(args.xml_file)
    create_nmap_report(scan_data, args.output_doc)
    print(f"Report saved as {args.output_doc}")

if __name__ == "__main__":
    main()
